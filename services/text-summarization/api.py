"""
Text Summarization Service API
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from pydantic import BaseModel
from typing import Optional
import os
import tempfile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document
from langchain_community.llms import Ollama
import pypdf
import docx
from io import BytesIO

router = APIRouter(prefix="/summarize", tags=["text-summarization"])

class SummarizeRequest(BaseModel):
    text: str
    summary_type: Optional[str] = "concise"  # concise, detailed, bullet_points
    max_length: Optional[int] = 200

class SummarizeResponse(BaseModel):
    success: bool
    summary: str
    original_length: int
    summary_length: int
    compression_ratio: float

def get_llm():
    """Get LLM instance - you can configure this based on your preferences"""
    try:
        # Try to use Ollama if available
        return Ollama(model="llama2", temperature=0.1)
    except Exception:
        # Fallback to a simple rule-based summarizer
        return None

def simple_extractive_summary(text: str, max_sentences: int = 3) -> str:
    """
    Simple extractive summarization when LLM is not available
    """
    sentences = text.split('. ')
    if len(sentences) <= max_sentences:
        return text
    
    # Simple scoring: prefer sentences with common words
    word_freq = {}
    words = text.lower().split()
    for word in words:
        word_freq[word] = word_freq.get(word, 0) + 1
    
    sentence_scores = []
    for i, sentence in enumerate(sentences):
        score = sum(word_freq.get(word.lower(), 0) for word in sentence.split())
        sentence_scores.append((score, i, sentence))
    
    # Get top sentences
    sentence_scores.sort(reverse=True)
    top_sentences = sorted(sentence_scores[:max_sentences], key=lambda x: x[1])
    
    return '. '.join([sentence[2] for sentence in top_sentences]) + '.'

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    pdf_file = BytesIO(file_content)
    pdf_reader = pypdf.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    docx_file = BytesIO(file_content)
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

@router.post("/text", response_model=SummarizeResponse)
async def summarize_text(request: SummarizeRequest):
    """
    Summarize plain text
    """
    try:
        llm = get_llm()
        
        if llm:
            # Use LangChain for summarization
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
            docs = [Document(page_content=request.text)]
            texts = text_splitter.split_documents(docs)
            
            chain = load_summarize_chain(llm, chain_type="map_reduce")
            summary = chain.run(texts)
        else:
            # Fallback to simple summarization
            max_sentences = 3 if request.summary_type == "concise" else 5
            if request.summary_type == "bullet_points":
                max_sentences = 4
            summary = simple_extractive_summary(request.text, max_sentences)
            
            if request.summary_type == "bullet_points":
                sentences = summary.split('. ')
                summary = '\n'.join([f"• {sentence.strip()}" for sentence in sentences if sentence.strip()])
        
        original_length = len(request.text)
        summary_length = len(summary)
        compression_ratio = summary_length / original_length if original_length > 0 else 0
        
        return SummarizeResponse(
            success=True,
            summary=summary,
            original_length=original_length,
            summary_length=summary_length,
            compression_ratio=compression_ratio
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Summarization failed: {str(e)}")

@router.post("/document", response_model=SummarizeResponse)
async def summarize_document(
    file: UploadFile = File(...),
    summary_type: str = "concise",
    max_length: int = 200
):
    """
    Summarize uploaded document (PDF, DOCX, TXT)
    """
    try:
        # Read file content
        file_content = await file.read()
        
        # Extract text based on file type
        if file.filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
        elif file.filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file_content)
        elif file.filename.lower().endswith('.txt'):
            text = file_content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or TXT files.")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in the document.")
        
        # Create request object and summarize
        request = SummarizeRequest(
            text=text,
            summary_type=summary_type,
            max_length=max_length
        )
        
        return await summarize_text(request)
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document processing failed: {str(e)}")

@router.get("/health")
async def health_check():
    """
    Health check endpoint
    """
    return {"status": "healthy", "service": "text-summarization"}

@router.get("/")
async def service_info():
    """
    Service information
    """
    return {
        "service": "Text Summarization Service",
        "description": "Summarize text content and documents using AI",
        "supported_formats": ["txt", "pdf", "docx"],
        "summary_types": ["concise", "detailed", "bullet_points"],
        "endpoints": [
            "/summarize/text - Summarize plain text",
            "/summarize/document - Summarize uploaded document",
            "/summarize/health - Health check"
        ]
    }