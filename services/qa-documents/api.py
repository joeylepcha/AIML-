"""
Q&A over Documents Service API
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from pydantic import BaseModel
from typing import List, Optional, Dict
import os
import tempfile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_community.llms import Ollama
from langchain.docstore.document import Document
import pypdf
import docx
from io import BytesIO
import uuid

router = APIRouter(prefix="/qa", tags=["qa-documents"])

# Global storage for document stores
document_stores: Dict[str, Chroma] = {}

class DocumentUploadResponse(BaseModel):
    success: bool
    message: str
    document_id: str
    filename: str
    pages_processed: int

class QuestionRequest(BaseModel):
    document_id: str
    question: str
    max_results: Optional[int] = 3

class QuestionResponse(BaseModel):
    success: bool
    answer: str
    confidence: float
    source_documents: List[Dict]
    document_id: str

def get_llm():
    """Get LLM instance"""
    try:
        return Ollama(model="llama2", temperature=0.1)
    except Exception:
        return None

def get_embeddings():
    """Get embeddings model"""
    try:
        return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    except Exception:
        # Fallback to a simple embedding
        return None

def extract_text_from_pdf(file_content: bytes) -> List[Document]:
    """Extract text from PDF file and return as documents"""
    pdf_file = BytesIO(file_content)
    pdf_reader = pypdf.PdfReader(pdf_file)
    documents = []
    
    for i, page in enumerate(pdf_reader.pages):
        text = page.extract_text()
        if text.strip():
            doc = Document(
                page_content=text,
                metadata={"page": i + 1, "source": "pdf"}
            )
            documents.append(doc)
    
    return documents

def extract_text_from_docx(file_content: bytes) -> List[Document]:
    """Extract text from DOCX file and return as documents"""
    docx_file = BytesIO(file_content)
    doc = docx.Document(docx_file)
    documents = []
    
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Split into chunks for better retrieval
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    text = "\n".join(full_text)
    chunks = text_splitter.split_text(text)
    
    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={"chunk": i + 1, "source": "docx"}
        )
        documents.append(doc)
    
    return documents

def extract_text_from_txt(file_content: bytes) -> List[Document]:
    """Extract text from TXT file and return as documents"""
    text = file_content.decode('utf-8')
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    chunks = text_splitter.split_text(text)
    documents = []
    
    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={"chunk": i + 1, "source": "txt"}
        )
        documents.append(doc)
    
    return documents

def simple_keyword_search(documents: List[Document], question: str, max_results: int = 3) -> List[Document]:
    """
    Simple keyword-based search when vector search is not available
    """
    question_words = set(question.lower().split())
    
    scored_docs = []
    for doc in documents:
        content_words = set(doc.page_content.lower().split())
        # Simple word overlap scoring
        overlap = len(question_words.intersection(content_words))
        if overlap > 0:
            scored_docs.append((overlap, doc))
    
    # Sort by score and return top results
    scored_docs.sort(key=lambda x: x[0], reverse=True)
    return [doc for _, doc in scored_docs[:max_results]]

def generate_simple_answer(relevant_docs: List[Document], question: str) -> str:
    """
    Generate a simple answer when LLM is not available
    """
    if not relevant_docs:
        return "I couldn't find relevant information to answer your question."
    
    # Combine relevant content
    context = "\n\n".join([doc.page_content[:500] for doc in relevant_docs])
    
    # Simple template-based response
    return f"Based on the document content, here's what I found:\n\n{context[:800]}..."

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload and process a document for Q&A
    """
    try:
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        
        # Read file content
        file_content = await file.read()
        
        # Extract text based on file type
        documents = []
        if file.filename.lower().endswith('.pdf'):
            documents = extract_text_from_pdf(file_content)
        elif file.filename.lower().endswith('.docx'):
            documents = extract_text_from_docx(file_content)
        elif file.filename.lower().endswith('.txt'):
            documents = extract_text_from_txt(file_content)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or TXT files.")
        
        if not documents:
            raise HTTPException(status_code=400, detail="No text content found in the document.")
        
        # Create vector store
        embeddings = get_embeddings()
        if embeddings:
            vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=embeddings,
                persist_directory=f"./chroma_db_{document_id}"
            )
            document_stores[document_id] = vectorstore
        else:
            # Store documents directly for keyword search fallback
            document_stores[document_id] = documents
        
        return DocumentUploadResponse(
            success=True,
            message="Document uploaded and processed successfully",
            document_id=document_id,
            filename=file.filename,
            pages_processed=len(documents)
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document upload failed: {str(e)}")

@router.post("/ask", response_model=QuestionResponse)
async def ask_question(request: QuestionRequest):
    """
    Ask a question about an uploaded document
    """
    try:
        if request.document_id not in document_stores:
            raise HTTPException(status_code=404, detail="Document not found. Please upload the document first.")
        
        store = document_stores[request.document_id]
        llm = get_llm()
        
        if isinstance(store, Chroma) and llm:
            # Use vector search with LLM
            qa = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=store.as_retriever(search_kwargs={"k": request.max_results}),
                return_source_documents=True
            )
            
            result = qa({"query": request.question})
            answer = result["result"]
            source_docs = [
                {
                    "content": doc.page_content[:200] + "...",
                    "metadata": doc.metadata
                }
                for doc in result["source_documents"]
            ]
            confidence = 0.8  # Placeholder confidence score
        
        else:
            # Fallback to simple keyword search
            documents = store if isinstance(store, list) else list(store.get())
            relevant_docs = simple_keyword_search(documents, request.question, request.max_results)
            answer = generate_simple_answer(relevant_docs, request.question)
            source_docs = [
                {
                    "content": doc.page_content[:200] + "...",
                    "metadata": doc.metadata
                }
                for doc in relevant_docs
            ]
            confidence = 0.6  # Lower confidence for simple search
        
        return QuestionResponse(
            success=True,
            answer=answer,
            confidence=confidence,
            source_documents=source_docs,
            document_id=request.document_id
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Question answering failed: {str(e)}")

@router.get("/documents")
async def list_documents():
    """
    List all uploaded documents
    """
    return {
        "success": True,
        "documents": list(document_stores.keys()),
        "count": len(document_stores)
    }

@router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """
    Delete an uploaded document
    """
    if document_id not in document_stores:
        raise HTTPException(status_code=404, detail="Document not found")
    
    del document_stores[document_id]
    
    # Clean up vector store files if they exist
    try:
        import shutil
        chroma_dir = f"./chroma_db_{document_id}"
        if os.path.exists(chroma_dir):
            shutil.rmtree(chroma_dir)
    except Exception:
        pass  # Ignore cleanup errors
    
    return {"success": True, "message": "Document deleted successfully"}

@router.get("/health")
async def health_check():
    """
    Health check endpoint
    """
    return {"status": "healthy", "service": "qa-documents"}

@router.get("/")
async def service_info():
    """
    Service information
    """
    return {
        "service": "Q&A over Documents Service",
        "description": "Upload documents and ask questions about their content",
        "supported_formats": ["txt", "pdf", "docx"],
        "endpoints": [
            "/qa/upload - Upload document for Q&A",
            "/qa/ask - Ask question about uploaded document",
            "/qa/documents - List uploaded documents",
            "/qa/documents/{id} - Delete document",
            "/qa/health - Health check"
        ]
    }